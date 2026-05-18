"""
pipeline/parser.py — extract plain text from resume files.

Supported formats:
  • PDF  — via PyMuPDF (fitz)
  • DOCX — via python-docx
  • DOC  — via antiword (system binary, installed in Docker)
"""
from __future__ import annotations

import io
import logging
import os
import re
import subprocess
import tempfile

logger = logging.getLogger(__name__)

# MIME types we accept
SUPPORTED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # .docx
    "application/msword",  # .doc
}

# Telegram file extensions we accept (fallback when mime_type is generic)
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}


def is_supported(filename: str, mime_type: str | None = None) -> bool:
    """Return True when the file can be parsed."""
    ext = os.path.splitext(filename.lower())[1]
    if ext in SUPPORTED_EXTENSIONS:
        return True
    if mime_type and mime_type in SUPPORTED_MIME_TYPES:
        return True
    return False


def extract_text(file_bytes: bytes, filename: str, mime_type: str | None = None) -> str:
    """
    Dispatch to the correct extractor based on filename extension / mime type.
    Returns clean plain text.
    Raises ValueError for unsupported formats.
    """
    ext = os.path.splitext(filename.lower())[1]

    if ext == ".pdf" or mime_type == "application/pdf":
        text = _extract_pdf(file_bytes)
    elif ext == ".docx" or mime_type == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        text = _extract_docx(file_bytes)
    elif ext == ".doc" or mime_type == "application/msword":
        text = _extract_doc(file_bytes)
    else:
        raise ValueError(
            f"Unsupported file format: extension='{ext}', mime='{mime_type}'. "
            f"Please upload a PDF, DOC, or DOCX file."
        )

    return _clean_text(text)


# ─────────────────────────────────────────────────────────────────────────────
# Private extractors
# ─────────────────────────────────────────────────────────────────────────────

def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:
        raise RuntimeError(
            "PyMuPDF is not installed. Add 'pymupdf' to pyproject.toml."
        ) from exc

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages: list[str] = []
    for page in doc:
        pages.append(page.get_text("text"))
    doc.close()

    text = "\n".join(pages)
    if not text.strip():
        raise ValueError(
            "Could not extract text from the PDF. "
            "It may be a scanned image. Please upload a text-based PDF."
        )
    return text


def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is not installed. Add 'python-docx' to pyproject.toml."
        ) from exc

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n".join(paragraphs)


def _extract_doc(file_bytes: bytes) -> str:
    """
    Extract text from legacy binary .doc files using the `antiword` system binary.

    antiword must be installed in the container:
        apt-get install -y antiword

    Falls back to attempting python-docx in case the file is actually .docx
    with a wrong extension.
    """
    # Attempt antiword first
    antiword_path = _find_binary("antiword")
    if antiword_path:
        return _run_antiword(file_bytes, antiword_path)

    # Fallback: try treating it as docx (some .doc files are actually newer format)
    logger.warning(
        "antiword not found; attempting python-docx fallback for .doc file."
    )
    try:
        return _extract_docx(file_bytes)
    except Exception:
        raise RuntimeError(
            "Could not parse the .doc file. "
            "Please convert it to DOCX or PDF and re-upload."
        )


def _run_antiword(file_bytes: bytes, antiword_path: str) -> str:
    """Write bytes to a temp file and run antiword."""
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        result = subprocess.run(
            [antiword_path, tmp_path],
            capture_output=True,
            text=True,
            timeout=30,
            encoding="utf-8",
            errors="replace",
        )
        if result.returncode != 0 or not result.stdout.strip():
            raise RuntimeError(
                f"antiword failed (rc={result.returncode}): {result.stderr.strip()}"
            )
        return result.stdout
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


def _find_binary(name: str) -> str | None:
    """Return full path to a system binary, or None if not found."""
    try:
        result = subprocess.run(
            ["which", name], capture_output=True, text=True, timeout=5
        )
        path = result.stdout.strip()
        return path if path else None
    except Exception:
        return None


# ─────────────────────────────────────────────────────────────────────────────
# Text cleaning
# ─────────────────────────────────────────────────────────────────────────────

def _clean_text(text: str) -> str:
    """Normalise whitespace and remove junk characters."""
    # Collapse runs of blank lines to a single blank line
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Strip trailing spaces on each line
    text = "\n".join(line.rstrip() for line in text.splitlines())
    # Remove non-printable control characters except newlines/tabs
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\x80-\xFF]", "", text)
    return text.strip()
